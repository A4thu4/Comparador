import streamlit as st
import pandas as pd
import numpy as np
import re
# Comparar Textos
from difflib import Differ, SequenceMatcher 
# Comparar Documentos
from PyPDF2 import PdfReader
import pdfplumber
import docx
import chardet 

# Função para comparar textos
def compare_texts(text1, text2):
    # Divide os textos em linhas
    text1_lines = [line.strip() for line in text1.splitlines() if line.strip()]
    text2_lines = [line.strip() for line in text2.splitlines() if line.strip()]
    
    # Usa SequenceMatcher para alinhar as linhas
    matcher = SequenceMatcher(None, text1_lines, text2_lines)
    
    # Prepara o resultado em HTML
    result = []
    result.append("""
    <style>
        .diff-container {
            display: flex;
            width: 100%;
            font-family: Calibri, Arial, sans-serif;
            border: 1px solid #ddd;
            border-radius: 6px;
            overflow: hidden;
            max-height: 80vh; 
            overflow-y: auto;
            overflow-x: auto;
        }
        .diff-column {
            width: 50%;
            min-width: 0;
            max-width: 50%;
            padding: 10px;
            margin: 0;
            border-right: 1px solid #eee;
            background: #fff;
            box-sizing: border-box;
            display: flex;
            flex-direction: column;
        }
        .diff-column:last-child {
            border-right: none;
        }
        .diff-line {
            white-space: pre-wrap;
            margin: 2px 0;
            padding: 2px;
            border-left: 4px solid transparent;
        }
        .diff-cell {
            flex: 1;
            padding: 2px 10px;
            white-space: pre-wrap;
        }
        .unchanged {
            background-color: #f8f8f8;
        }
        .deleted {
            background-color: #ffdddd;
            text-decoration: line-through;
            border-left: 4px solid #ff6f6f;
        }
        .added {
            background-color: #ddffdd;
            text-decoration: underline;
            border-left: 4px solid #4fe87b;
        }
        .changed-old {
            background-color: #ff758f;
            text-decoration: line-through;
        }
        .changed-new {
            background-color: #abff4f;
            text-decoration: underline;
        }
        .line-number {
            color: #999;
            margin-right: 5px;
            user-select: none;
            width: 20px;
            display: inline-block;
        }
    </style>
    <div class="diff-container">
        <div class="diff-column">
            <h3>Texto Original</h3>
    """)
    
    left_lines = []
    right_lines = []
    
    # Processa cada bloco de diferenças
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():

        if tag == 'equal':
            for line in text1_lines[i1:i2]:
                left_lines.append(('unchanged', line))
                right_lines.append(('unchanged', line))

        elif tag == 'delete':
            for line in text1_lines[i1:i2]:
                left_lines.append(('deleted', line))
                right_lines.append(('empty', ''))

        elif tag == 'insert':
            for line in text2_lines[j1:j2]:
                if left_lines and left_lines[-1][0] == 'empty':
                    right_lines[-1] = ('added', line)
                else:
                    left_lines.append(('empty',''))
                    right_lines.append(('added',line))

        elif tag == 'replace':
            matched_pairs = []
            used_new = set()
            used_old = set()
            # Pareamento linha a linha
            for i, old_line in enumerate(text1_lines[i1:i2]):
                best_match = None
                best_ratio = 0.8
                for j, new_line in enumerate(text2_lines[j1:j2]):
                    if j in used_new:
                        continue
                    ratio = SequenceMatcher(None, old_line, new_line).ratio()
                    if ratio > best_ratio:
                        best_ratio = ratio
                        best_match = j
                if best_match is not None:
                    matched_pairs.append((i, best_match))
                    used_new.add(best_match)
                    used_old.add(i)

            # Agora percorre ambos os blocos na ordem original
            old_idx, new_idx = 0, 0
            len_old = i2 - i1
            len_new = j2 - j1
            while old_idx < len_old or new_idx < len_new:
                # Se ambos são pareados
                pair = next(((oi, nj) for oi, nj in matched_pairs if oi == old_idx and nj == new_idx), None)
                if pair:
                    old_line = text1_lines[i1 + old_idx]
                    new_line = text2_lines[j1 + new_idx]
                    d = Differ()
                    diff = list(d.compare(old_line.split(), new_line.split()))
                    old_text = []
                    new_text = []
                    for word in diff:
                        if word.startswith('- '):
                            old_text.append(f'<span class="changed-old">{word[2:]}</span>')
                        elif word.startswith('+ '):
                            new_text.append(f'<span class="changed-new">{word[2:]}</span>')
                        elif word.startswith('  '):
                            old_text.append(word[2:])
                            new_text.append(word[2:])
                    left_lines.append(('changed', ' '.join(old_text)))
                    right_lines.append(('changed', ' '.join(new_text)))
                    old_idx += 1
                    new_idx += 1
                elif old_idx < len_old and old_idx not in used_old:
                    left_lines.append(('deleted', text1_lines[i1 + old_idx]))
                    right_lines.append(('empty', ''))
                    old_idx += 1
                elif new_idx < len_new and new_idx not in used_new:
                    left_lines.append(('empty', ''))
                    right_lines.append(('added', text2_lines[j1 + new_idx]))
                    new_idx += 1
                else:
                    old_idx += 1
                    new_idx += 1

    # Adiciona as linhas do lado esquerdo (original)
    for i, (line_class, line) in enumerate(left_lines):
        if line_class == 'empty':
            result.append(f'<div class="diff-line empty"><span class="line-number">{i+1}</span>&nbsp;</div>')
        else:
            result.append(f'<div class="diff-line {line_class}"><span class="line-number">{i+1}</span>{line}</div>')
    
    result.append("""
        </div>
        <div class="diff-column">
            <h3>Texto Modificado</h3>
    """)
    
    # Adiciona as linhas do lado direito (modificado)
    for i, (line_class, line) in enumerate(right_lines):
        if line_class == 'empty':
            result.append(f'<div class="diff-line empty"><span class="line-number">{i+1}</span>&nbsp;</div>')
        else:
            result.append(f'<div class="diff-line {line_class}"><span class="line-number">{i+1}</span>{line}</div>')
    
    result.append("""
        </div>
    </div>
    """)
    
    return ''.join(result)

# Funções para comparar arquivos 
def is_footer_row(row):
    """Retorna True se a linha parece ser rodapé."""
    if not row or all((cell is None or str(cell).strip() == "") for cell in row):
        return True  # Linha vazia
    row_str = " ".join([str(cell) for cell in row if cell is not None]).strip()
    # Ajuste os padrões conforme necessário para seu caso
    footer_patterns = [
        r"portaria sead", r"sei \d+", r"p[áa]g\.\s*\d+", r"n[úu]mero", r"processo", r"^none$", r"^$", r"^página",
        r"^.*documento.*$", r"^.*secretaria.*$", r"^.*governo.*$", r"^.*cnpj.*$"
    ]
    for pat in footer_patterns:
        if re.search(pat, row_str, re.IGNORECASE):
            return True
    
    return False

def extract_tables_from_pdf(file):
    def make_unique(seq):
        seen = {}
        result = []
        for item in seq:
            if item not in seen:
                seen[item] = 0
                result.append(item)
            else:
                seen[item] += 1
                result.append(f"{item}_{seen[item]}")
        return result

    full_table = []
    header = None
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            tables_on_page = page.extract_tables()
            for table in tables_on_page:
                if table:
                    # Só remove rodapé se estiver na última linha
                    filtered_table = table.copy()
                    while filtered_table and is_footer_row(filtered_table[-1]):
                        filtered_table = filtered_table[:-1]
                    if not filtered_table:
                        continue
                    if not header:
                        header = filtered_table[0]
                        full_table.extend(filtered_table[1:])
                    else:
                        if filtered_table[0] == header:
                            full_table.extend(filtered_table[1:])
                        else:
                            full_table.extend(filtered_table)
    if header and full_table:
        unique_header = make_unique(header)
        return [pd.DataFrame(full_table, columns=unique_header)]
    else:
        return []

def extract_tables_from_docx(file):
    doc = docx.Document(file)
    tables = []
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        df = pd.DataFrame(data)
        tables.append(df)
    return tables

def extract_text(file):
    """Extrai e padroniza texto de diferentes formatos de arquivo"""
    if not file:
        return None
        
    file_extension = file.name.split('.')[-1].lower()
    if file_extension not in ['pdf', 'docx', 'doc', 'txt', 'csv']:
        st.error(f"Formato não suportado: {file_extension}")
        return None
    
    try:
        # Padroniza o tratamento de encoding para todos os formatos
        def decode_text(raw_data):
            """Função auxiliar para decodificar texto com detecção de encoding"""
            if not raw_data:
                return ""
                
            # Detecta encoding com confiança mínima de 70%
            detected = chardet.detect(raw_data)
            encoding = detected['encoding'] if detected['confidence'] > 0.7 else 'utf-8'
            
            try:
                return raw_data.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                # Tenta utf-8 com fallback para substituição de caracteres inválidos
                return raw_data.decode('utf-8', errors='replace')
        
        # PDF - Mantém estrutura original
        if file_extension == 'pdf':
            try:
                reader = PdfReader(file)
                num_paginas = len(reader.pages)
                texto_completo = []
                for i in range(num_paginas):
                    pagina = reader.pages[i]
                    texto_pagina = pagina.extract_text()
                    texto_completo.append(texto_pagina)
                texto_final = '\n'.join(texto_completo) if texto_completo else None
                if texto_final:
                    texto_final = texto_final.replace('\r\n', '\n').replace('\r', '\n')
                    linhas = texto_final.split('\n')
                    paragrafos = []
                    paragrafo_atual = []
                    for linha in linhas:
                        # Remove numeração do início da linha (ex: "1. ", "2. ", "10. ")
                        linha_sem_num = linha.lstrip()
                        import re
                        linha_sem_num = re.sub(r'^\d+\.\s*', '', linha_sem_num)
                        if linha_sem_num.strip() == "":
                            if paragrafo_atual:
                                paragrafos.append(' '.join(paragrafo_atual).strip())
                                paragrafo_atual = []
                        else:
                            paragrafo_atual.append(linha_sem_num.strip())
                    if paragrafo_atual:
                        paragrafos.append(' '.join(paragrafo_atual).strip())
                    texto_final = '\n\n'.join(paragrafos)
                return texto_final if texto_final else None
            except Exception as e:
                print(f"Erro ao ler o PDF: {e}")
                return None
            
        # DOCX - Padroniza espaçamento entre parágrafos
        elif file_extension in ['docx', 'doc']:
            try:
                doc = docx.Document(file)
                text_lines = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_lines.append(para.text.strip())
                    else:
                        text_lines.append('')
                
                text = '\n'.join(text_lines)
                return text.strip() if text.strip() else None
            except Exception as e:
                st.error(f"Erro ao ler arquivo Word: {str(e)}")
                return None

        # TXT/CSV - Padroniza tratamento de quebras de linha
        elif file_extension in ['txt', 'csv']:
            file.seek(0)
            raw_data = file.read()
            text = decode_text(raw_data)
            
            if not text:
                st.warning("O arquivo está vazio.")
                return None
                
            # Padroniza quebras de linha para \n
            text = text.replace('\r\n', '\n').replace('\r', '\n')
            # Remove numeração do início de cada linha
            import re
            linhas = text.split('\n')
            linhas_sem_enum = [
                            re.sub(
                                r'^\s*((\d+(\.\d+)*[.)]?)|[.\-•])[\s\t\u00A0]*',  # cobre . espaço, .\t, . , 1.2.3. etc
                                '',
                                linha
                            ).strip()
                            for linha in linhas
                        ]            
            texto_final = '\n'.join(linhas_sem_enum)

            # Para CSV, trata como texto puro (não tenta parsear)
            return texto_final.strip() if texto_final.strip() else None

        else:
            st.error(f"Formato não suportado: {file_extension}")
            return None
            
    except Exception as e:
        st.error(f"Erro ao processar arquivo: {str(e)}")
        return None

def compare_docs(doc1, doc2):
    """Compara dois documentos com visualização lado a lado"""
    text1 = extract_text(doc1) or ""
    text2 = extract_text(doc2) or ""

    if not text1 and not text2:
        return None, None, True
    
    if text1.strip() == text2.strip():
        return None, None, True
    
    # Divide os textos em linhas mantendo as quebras originais
    text1_lines = text1.splitlines()
    text2_lines = text2.splitlines()
    
    # Remove linhas vazias do início e fim, mas mantém as do meio
    text1_lines = [line.rstrip() for line in text1_lines]
    text2_lines = [line.rstrip() for line in text2_lines]
    
    # Usa SequenceMatcher para alinhar as linhas
    matcher = SequenceMatcher(None, text1_lines, text2_lines)
    
    # Prepara o resultado em HTML 
    result = []
    result.append("""
    <style>
        .diff-container {
            display: flex;
            width: 100%;
            font-family: Calibri, Arial, sans-serif;
            border: 1px solid #ddd;
            border-radius: 5px;
            overflow: hidden;
            max-height: 80vh; 
            overflow-y: auto;
            overflow-x: auto;
        }
        .diff-column {
            width: 50%;
            min-width: 0;
            max-width: 50%;
            padding: 10px;
            margin: 0;
            border-right: 1px solid #eee;
            background: #fff;
            box-sizing: border-box;
            display: flex;
            flex-direction: column;
        }
        .diff-column:last-child {
            border-right: none;
        }
        .diff-line {
            white-space: pre-wrap;
            margin: 2px 0;
            padding: 2px;
            border-left: 4px solid transparent;
        }
        .unchanged {
            background-color: #f8f8f8;
        }
        .deleted {
            background-color: #ffdddd;
            text-decoration: line-through;
            border-left: 4px solid #ff6f6f;
        }
        .added {
            background-color: #ddffdd;
            text-decoration: underline;
            border-left: 4px solid #4fe87b;
        }
        .changed-old {
            background-color: #ff6f6f;
            text-decoration: line-through;
        }
        .changed-new {
            background-color: #4fe87b;
            text-decoration: underline;
        }
        .line-number {
            color: #999;
            margin-right: 5px;
            user-select: none;
            width: 20px;
            display: inline-block;
        }
        .diff-header {
            background: #f5f5f5;
            padding: 10px;
            font-weight: bold;
            border-bottom: 1px solid #ddd;
            margin: -10px -10px 10px -10px;
        }
        .empty-line {
            height: 5px;
            margin: 2px 0;
            background: repeating-linear-gradient(
                45deg,
                #f0f0f0,
                #f0f0f0 2px,
                white 2px,
                white 4px
            );
        }
    </style>
    <div class="diff-container">
        <div class="diff-column">
            <div class="diff-header">Documento Original</div>
    """)
    
    left_lines = []
    right_lines = []
    changes_report = []
    line_counter = 1
    
    # Processa cada bloco de diferenças
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            for line in text1_lines[i1:i2]:
                if line.strip() == "":
                    left_lines.append(('empty-line', '', line_counter))
                    right_lines.append(('empty-line', '', line_counter))
                else:
                    left_lines.append(('unchanged', line, line_counter))
                    right_lines.append(('unchanged', line, line_counter))
                line_counter += 1

        elif tag == 'delete':
            for line in text1_lines[i1:i2]:
                if line.strip() == "":
                    left_lines.append(('empty-line', '', line_counter))
                    right_lines.append(('empty-line', '', line_counter))
                else:
                    left_lines.append(('deleted', line, line_counter))
                    right_lines.append(('empty', '', line_counter))
                    changes_report.append({
                        'Tipo': 'Removido',
                        'Conteúdo': line,
                        'Localização': f'Linha {line_counter}'
                    })
                line_counter += 1

        elif tag == 'insert':
            for line in text2_lines[j1:j2]:
                if line.strip() == "":
                    left_lines.append(('empty-line', '', line_counter))
                    right_lines.append(('empty-line', '', line_counter))
                else:
                    if left_lines and left_lines[-1][0] == 'empty':
                        right_lines[-1] = ('added', line, right_lines[-1][2])
                    else:
                        left_lines.append(('empty', '', line_counter))
                        right_lines.append(('added', line, line_counter))
                    changes_report.append({
                        'Tipo': 'Adicionado',
                        'Conteúdo': line,
                        'Localização': f'Linha {line_counter}'
                    })
                line_counter += 1

        elif tag == 'replace':
            # Processa substituições mantendo a estrutura de linhas
            max_lines = max((i2-i1), (j2-j1))
            for n in range(max_lines):
                old_line = text1_lines[i1 + n] if n < (i2-i1) else ""
                new_line = text2_lines[j1 + n] if n < (j2-j1) else ""
                
                if old_line.strip() == "" and new_line.strip() == "":
                    left_lines.append(('empty-line', '', line_counter))
                    right_lines.append(('empty-line', '', line_counter))
                elif old_line.strip() == "":
                    left_lines.append(('empty', '', line_counter))
                    right_lines.append(('added', new_line, line_counter))
                    changes_report.append({
                        'Tipo': 'Adicionado',
                        'Conteúdo': new_line,
                        'Localização': f'Linha {line_counter}'
                    })
                elif new_line.strip() == "":
                    left_lines.append(('deleted', old_line, line_counter))
                    right_lines.append(('empty', '', line_counter))
                    changes_report.append({
                        'Tipo': 'Removido',
                        'Conteúdo': old_line,
                        'Localização': f'Linha {line_counter}'
                    })
                else:
                    # Comparação detalhada dentro da linha
                    d = Differ()
                    diff = list(d.compare(old_line.split(), new_line.split()))
                    
                    old_text = []
                    new_text = []
                    
                    for word in diff:
                        if word.startswith('- '):
                            old_text.append(f'<span class="changed-old">{word[2:]}</span>')
                        elif word.startswith('+ '):
                            new_text.append(f'<span class="changed-new">{word[2:]}</span>')
                        elif word.startswith('  '):
                            old_text.append(word[2:])
                            new_text.append(word[2:])
                    
                    left_lines.append(('changed', ' '.join(old_text), line_counter))
                    right_lines.append(('changed', ' '.join(new_text), line_counter))
                    
                    changes_report.append({
                        'Tipo': 'Alterado',
                        'Original': old_line,
                        'Modificado': new_line,
                        'Localização': f'Linha {line_counter}'
                    })
                
                line_counter += 1

    # Adiciona as linhas do lado esquerdo (original)
    for line_class, line, line_num in left_lines:
        if line_class == 'empty-line':
            result.append(f'<div class="empty-line" title="Linha {line_num}"></div>')
        elif line_class == 'empty':
            result.append(f'<div class="diff-line empty"><span class="line-number">{line_num}</span>&nbsp;</div>')
        else:
            result.append(f'<div class="diff-line {line_class}"><span class="line-number">{line_num}</span>{line}</div>')

    result.append("""
        </div>
        <div class="diff-column">
            <div class="diff-header">Documento Modificado</div>
    """)
    
    # Adiciona as linhas do lado direito (modificado)
    for line_class, line, line_num in right_lines:
        if line_class == 'empty-line':
            result.append(f'<div class="empty-line" title="Linha {line_num}"></div>')
        elif line_class == 'empty':
            result.append(f'<div class="diff-line empty"><span class="line-number">{line_num}</span>&nbsp;</div>')
        else:
            result.append(f'<div class="diff-line {line_class}"><span class="line-number">{line_num}</span>{line}</div>')
    
    result.append("""
        </div>
    </div>
    """)
    
    # Gerar relatório de alterações
    diff_df = pd.DataFrame(changes_report) if changes_report else None
    
    return ''.join(result), diff_df, False

def compare_table_list(tables1, tables2):
    results = []
    for i in range(max(len(tables1), len(tables2))):
        df1 = tables1[i] if i < len(tables1) else pd.DataFrame()
        df2 = tables2[i] if i < len(tables2) else pd.DataFrame()
        all_cols = list(df1.columns.union(df2.columns)) if not df1.empty and not df2.empty else []
        result = smart_row_comparison(df1, df2, all_cols)
        results.append(result)
    return results

# Funções para comparar planilhas Excel
def compare_excel(file1, file2, selected_sheet=None):
    try:
        xls1 = pd.ExcelFile(file1)
        xls2 = pd.ExcelFile(file2)
    except Exception as e:
        st.error(f"Erro ao ler arquivos: {e}")
        return None

    # Padroniza nomes das abas
    sheets1 = set(s.strip().lower() for s in xls1.sheet_names)
    sheets2 = set(s.strip().lower() for s in xls2.sheet_names)

    if sheets1 != sheets2:
        st.warning(f"Os arquivos não possuem as mesmas abas.\n")
        return None
    
    # Determinar quais abas comparar
    all_sheets = sorted(set(xls1.sheet_names) | set(xls2.sheet_names))
    
    if selected_sheet:
        if selected_sheet not in all_sheets:
            st.warning(f"A aba '{selected_sheet}' não foi encontrada em ambos arquivos.")
            return None
        sheets_to_compare = [selected_sheet]
    else:
        sheets_to_compare = all_sheets

    results = {}
    
    for sheet in sheets_to_compare:
        # Carregar dados
        df1 = xls1.parse(sheet) if sheet in xls1.sheet_names else pd.DataFrame()
        df2 = xls2.parse(sheet) if sheet in xls2.sheet_names else pd.DataFrame()

        # Remover índices e resetar
        df1 = df1.reset_index(drop=True)
        df2 = df2.reset_index(drop=True)
        
        # Garantir que as colunas sejam as mesmas
        all_cols = list(df1.columns.union(df2.columns))
        df1 = df1.reindex(columns=all_cols)
        df2 = df2.reindex(columns=all_cols)
        
        # Usar algoritmo inteligente de comparação de linhas
        comparison_result = smart_row_comparison(df1, df2, all_cols)
        
        results[sheet] = comparison_result
    
    return results if not selected_sheet else results.get(selected_sheet)

def smart_row_comparison(df1, df2, all_cols):
    """
    Algoritmo inteligente que detecta inserções, deleções e alterações reais
    sem marcar linhas deslocadas como alteradas
    """
    # Proteção contra dataframes sem colunas
    if not all_cols:
        all_cols = list(df1.columns) if not df1.empty else list(df2.columns)

    # Se ainda estiver vazio, retorna dataframes vazios e stats zeradas
    if not all_cols:
        empty_df = pd.DataFrame()
        return (empty_df, empty_df, {'changes': 0, 'additions': 0, 'deletions': 0, 'total_cells': 0})
    all_cols = pd.Index(all_cols).drop_duplicates().tolist()
    df1 = df1.loc[:, ~df1.columns.duplicated()]
    df2 = df2.loc[:, ~df2.columns.duplicated()]
    
    df1 = df1.reindex(columns=all_cols)
    df2 = df2.reindex(columns=all_cols)

    # Padroniza valores vazios para string vazia
    df1 = df1.replace(["None", "nan", None, np.nan], "", regex=True)
    df2 = df2.replace(["None", "nan", None, np.nan], "", regex=True)

    # Converter linhas para strings para comparação
    def row_to_string(row):
        return '|'.join([str(val) if (pd.notna(val) and str(val).strip().lower() not in ["none", "nan"]) else '' for val in row])
    
    # Criar listas de strings representando cada linha
    rows1 = [row_to_string(df1.iloc[i]) for i in range(len(df1))]
    rows2 = [row_to_string(df2.iloc[i]) for i in range(len(df2))]
    
    # Usar SequenceMatcher para detectar operações (igual, inserir, deletar, substituir)
    matcher = SequenceMatcher(None, rows1, rows2)
    
    # Preparar DataFrames de resultado
    max_rows = len(df1) + len(df2)
    result_df1 = pd.DataFrame(index=range(max_rows), columns=all_cols)
    result_df2 = pd.DataFrame(index=range(max_rows), columns=all_cols)
    style_df1 = pd.DataFrame('', index=range(max_rows), columns=all_cols)
    style_df2 = pd.DataFrame('', index=range(max_rows), columns=all_cols)
    
    def normalize(val):
        if pd.isna(val):
            return ''
        return str(val).strip().lower().replace('–', '-').replace('“', '"').replace('”', '"')
    
    # Contadores para estatísticas
    changes_count = 0
    additions_count = 0
    deletions_count = 0
    
    current_row = 0
    
    # Processar cada operação do SequenceMatcher
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            # Linhas iguais - copiar sem marcação
            for i, (idx1, idx2) in enumerate(zip(range(i1, i2), range(j1, j2))):
                row_idx = current_row + i
                result_df1.iloc[row_idx] = df1.iloc[idx1].values.tolist()
                result_df2.iloc[row_idx] = df2.iloc[idx2].values.tolist()
                # Sem estilo especial para linhas iguais
            current_row += (i2 - i1)
        
        elif tag == 'delete':
            # Linhas deletadas - aparecem apenas no arquivo 1
            for i, idx1 in enumerate(range(i1, i2)):
                row_idx = current_row + i
                result_df1.iloc[row_idx] = df1.iloc[idx1].values.tolist()
                # Linha vazia no arquivo 2
                for col in all_cols:
                    style_df1.iloc[row_idx, style_df1.columns.get_loc(col)] = "background-color: #f8d7da; border-left: 4px solid #dc3545;"
                    style_df2.iloc[row_idx, style_df2.columns.get_loc(col)] = "background-color: #e0e0e9; border-left: 4px solid #6c757d;"
                deletions_count += 1
            current_row += (i2 - i1)
            
        elif tag == 'insert':
            # Linhas inseridas - aparecem apenas no arquivo 2
            for i, idx2 in enumerate(range(j1, j2)):
                row_idx = current_row + i
                result_df2.iloc[row_idx] = df2.iloc[idx2].values.tolist()
                # Linha vazia no arquivo 1
                for col in all_cols:
                    style_df1.iloc[row_idx, style_df1.columns.get_loc(col)] = "background-color: #e0e0e9; border-left: 4px solid #6c757d;"
                    style_df2.iloc[row_idx, style_df2.columns.get_loc(col)] = "background-color: #d4edda; border-left: 4px solid #28a745;"
                additions_count += 1
            current_row += (j2 - j1)
            
        elif tag == 'replace':
            max_replace_rows = max(i2 - i1, j2 - j1)
            for i in range(max_replace_rows):
                row_idx = current_row + i

                # Se não existe linha correspondente, cria linha vazia
                if i < (i2 - i1):
                    row1 = df1.iloc[i1 + i]
                else:
                    row1 = pd.Series([""] * len(all_cols), index=all_cols)
                if i < (j2 - j1):
                    row2 = df2.iloc[j1 + i]
                else:
                    row2 = pd.Series([""] * len(all_cols), index=all_cols)

                # Se uma linha está vazia e a outra não, trata como adição ou remoção
                if row1.isnull().all() or (row1 == "").all():
                    # Linha só no arquivo 2 (adição)
                    result_df2.iloc[row_idx] = row2.values.tolist()
                    for col in all_cols:
                        style_df1.iloc[row_idx, style_df1.columns.get_loc(col)] = "background-color: #e0e0e9; border-left: 4px solid #6c757d;"
                        style_df2.iloc[row_idx, style_df2.columns.get_loc(col)] = "background-color: #d4edda; border-left: 4px solid #28a745;"
                    additions_count += 1
                elif row2.isnull().all() or (row2 == "").all():
                    # Linha só no arquivo 1 (remoção)
                    result_df1.iloc[row_idx] = row1.values.tolist()
                    for col in all_cols:
                        style_df1.iloc[row_idx, style_df1.columns.get_loc(col)] = "background-color: #f8d7da; border-left: 4px solid #dc3545;"
                        style_df2.iloc[row_idx, style_df2.columns.get_loc(col)] = "background-color: #e0e0e9; border-left: 4px solid #6c757d;"
                    deletions_count += 1
                else:
                    # Linhas diferentes: alteração célula a célula
                    result_df1.iloc[row_idx] = row1.values.tolist()
                    result_df2.iloc[row_idx] = row2.values.tolist()
                    for col in all_cols:
                        val1 = row1[col] if col in row1 else None
                        val2 = row2[col] if col in row2 else None
                        if pd.isna(val1) and pd.isna(val2):
                            continue
                        elif (str(val1).strip() == "" and str(val2).strip() == ""):
                            continue
                        elif str(val1).strip() != str(val2).strip():
                            style_df1.iloc[row_idx, style_df1.columns.get_loc(col)] = "background-color: #ffff99; border-left: 4px solid #ffc107;"
                            style_df2.iloc[row_idx, style_df2.columns.get_loc(col)] = "background-color: #ffff99; border-left: 4px solid #ffc107;"
                            changes_count += 1
            current_row += max_replace_rows

    result_df1 = result_df1.iloc[:current_row].copy()
    result_df2 = result_df2.iloc[:current_row].copy()
    style_df1 = style_df1.iloc[:current_row].copy()
    style_df2 = style_df2.iloc[:current_row].copy()

    # Padroniza valores vazios para string vazia
    result_df1 = result_df1.replace(["None", "nan", None, np.nan], "", regex=True)
    result_df2 = result_df2.replace(["None", "nan", None, np.nan], "", regex=True)

    # Aplicar estilos aos DataFrames
    styled_df1 = result_df1.style.apply(
        lambda col: style_df1[col.name] if col.name in style_df1.columns else [''] * len(result_df1), 
        axis=0
    )
    styled_df2 = result_df2.style.apply(
        lambda col: style_df2[col.name] if col.name in style_df2.columns else [''] * len(result_df2), 
        axis=0
    )
    
    # Calcular estatísticas
    stats = {
        'changes': changes_count,
        'additions': additions_count,
        'deletions': deletions_count,
        'total_cells': len(all_cols) * max_rows
    }
    
    return (styled_df1, styled_df2, stats)

def display_excel_comparison(result, sheet_name, file1_name, file2_name):
    """Exibe a comparação do Excel com estilo profissional tipo diffchecker"""
    if len(result) == 3:
        styled_df1, styled_df2, stats = result
    else:
        styled_df1, styled_df2 = result
        stats = {'changes': 0, 'additions': 0, 'deletions': 0, 'total_cells': 0}
    
    # Legenda
    st.markdown("""
    <div style="background-color: #F3F3F3; padding: 15px; border-radius: 8px; margin: 15px 0;">
        <h5 style="margin-top: 0; color: #333;">Legenda:</h5>
        <div style="display: flex; gap: 20px; flex-wrap: wrap;">
            <div style="display: flex; align-items: center; gap: 5px;">
                <div style="width: 20px; height: 20px; background-color: #ffff99; border-left: 4px solid #ffc107; border-radius: 3px;"></div>
                <span>Alterados</span>
            </div>
            <div style="display: flex; align-items: center; gap: 5px;">
                <div style="width: 20px; height: 20px; background-color: #d4edda; border-left: 4px solid #28a745; border-radius: 3px;"></div>
                <span>Adicionados</span>
            </div>
            <div style="display: flex; align-items: center; gap: 5px;">
                <div style="width: 20px; height: 20px; background-color: #f8d7da; border-left: 4px solid #dc3545; border-radius: 3px;"></div>
                <span>Removidos</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Comparação lado a lado
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown(f"""
        <div style="background-color: #F3F3F3; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
            <h4 style="margin: 0;">{file1_name}</h4>
            <small>Aba: {sheet_name}</small>
        </div>
        """, unsafe_allow_html=True)
        st.dataframe(styled_df1, use_container_width=True, height=600)
    # st.divider()
    with col2:
        st.markdown(f"""
        <div style="background-color: #F3F3F3; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
            <h4 style="margin: 0;">{file2_name}</h4>
            <small>Aba: {sheet_name}</small>
        </div>
        """, unsafe_allow_html=True)
        st.dataframe(styled_df2, use_container_width=True, height=600)
    
    # Botão de download
    if stats['changes'] + stats['additions'] + stats['deletions'] > 0:
        # Gerar relatório HTML para download
        html_report = generate_excel_report(styled_df1, styled_df2, stats, sheet_name, file1_name, file2_name)
        # st.download_button(
        #     label="Baixar Comparação Tabelas",
        #     data=html_report.encode("utf-8"),
        #     file_name=f"Comparacao_Excel_{sheet_name}.html",
        #     mime="text/html",
        #     type="secondary"
        # )

def generate_excel_report(styled_df1, styled_df2, stats, sheet_name, file1_name, file2_name):
    """Gera relatório HTML da comparação para download"""
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Relatório de Comparação Excel - {sheet_name}</title>
        <meta charset="utf-8">
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; }}
            .header {{ background-color: #f8f9fa; padding: 20px; border-radius: 8px; margin-bottom: 20px; }}
            .stats {{ display: flex; gap: 20px; margin: 15px 0; }}
            .stat {{ background: white; padding: 15px; border-radius: 5px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
            .legend {{ background-color: #f8f9fa; padding: 15px; border-radius: 8px; margin: 15px 0; }}
            .comparison {{ display: flex; gap: 20px; }}
            .file-section {{ flex: 1; }}
            .file-header {{ background-color: #e3f2fd; padding: 10px; border-radius: 5px; margin-bottom: 10px; }}
            table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #f5f5f5; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>📊 Relatório da Comparação </h1>
            <p><strong>Arquivos:</strong> {file1_name} // {file2_name}</p>
            <p><strong>Aba:</strong> {sheet_name}</p>
        </div>
        
        <div class="stats">
            <div class="stat">
                <h3>🔄 Alterações</h3>
                <p style="font-size: 24px; margin: 0;">{stats['changes']}</p>
            </div>
            <div class="stat">
                <h3>➕ Adições</h3>
                <p style="font-size: 24px; margin: 0;">{stats['additions']}</p>
            </div>
            <div class="stat">
                <h3>➖ Remoções</h3>
                <p style="font-size: 24px; margin: 0;">{stats['deletions']}</p>
            </div>
        </div>
        
        <div class="legend">
            <h3>Legenda das Cores:</h3>
            <p>• <span style="background-color: #fff3cd; padding: 2px 8px; border-left: 4px solid #ffc107;">Valores alterados</span></p>
            <p>• <span style="background-color: #d4edda; padding: 2px 8px; border-left: 4px solid #28a745;">Valores adicionados</span></p>
            <p>• <span style="background-color: #f8d7da; padding: 2px 8px; border-left: 4px solid #dc3545;">Valores removidos</span></p>
        </div>
        
        <div class="comparison">
            <div class="file-section">
                <div class="file-header">
                    <h3>{file1_name}</h3>
                </div>
                {styled_df1.to_html()}
            </div>
            
            <div class="file-section">
                <div class="file-header">
                    <h3>{file2_name}</h3>
                </div>
                {styled_df2.to_html()}
            </div>
        </div>
    </body>
    </html>
    """
    return html

def excel_equal(file1, file2):
    try:
        xls1 = pd.ExcelFile(file1)
        xls2 = pd.ExcelFile(file2)
    except Exception as e:
        return False

    sheets1 = set(xls1.sheet_names)
    sheets2 = set(xls2.sheet_names)
    if sheets1 != sheets2:
        return False

    for sheet in sheets1:
        df1 = xls1.parse(sheet)
        df2 = xls2.parse(sheet)
        # Converte nomes das colunas para string antes de ordenar
        df1.columns = df1.columns.map(str)
        df2.columns = df2.columns.map(str)
        df1 = df1.sort_index(axis=1).sort_index()
        df2 = df2.sort_index(axis=1).sort_index()
        if not df1.equals(df2):
            return False
    return True

def main():
# Configuração da página
    st.set_page_config(page_title="Comparador GNCP", page_icon="utils/Brasão.png", layout="wide")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.image("utils/Logomarca SEAD 2.png", width=800)
        
# CSS customizado 
    st.markdown(
        """
        <style>
            img {
                margin-top: -3rem !important;
                margin-bottom: -1.2rem !important;
                align: center !important;
            }
            h1 {
                font-size: 2.12rem !important;
                margin-bottom: 1rem !important;
                margin-left: 1.6rem !important;
            }
            :root {
                --primary-color: #1bb50b !important;  /* Verde */
                --background-color: #FFFFFF !important;  /* Branco */
                --secondary-background-color: #FFFFFF !important;  /* Branco */
                --text-color: #000000 !important;  /* Preto */
            }

            /* Aplica cinza SOMENTE nos inputs */
            .stTextInput>div>div>input,
            .stNumberInput>div>div>input,
            .stTextArea>div>div>textarea,
            .stSelectbox>div>div>select,
            .stDateInput>div>div>input {
                background-color: #F3F3F3 !important;  /* Cinza claro */
                border-radius: 8px !important;
            }

            /* Mantém fundo branco em outros containers */
            .stApp, .stSidebar, .stAlert, .stMarkdown {
                background-color: #FFFFFF !important;
            }

            /* Estilo para botões */
            .stButton > button {
                border-radius: 8px !important;
                border: 1px solid #e0e0e0 !important;
                transition: all 0.3s ease !important;
                font-weight: 500 !important;
                color: #ff666f !important;
            }

            .stButton > button:hover {
                background: linear-gradient(135deg, #FFF, #FFF) !important;
                color: #ff666f !important; /* texto verde */
                border: 2px solid #ff666f !important; /* borda verde */
                box-shadow: 0 2px 8px rgba(27,181,11,0.15) !important; /* sombra suave */
                transform: translateY(-2px) scale(1.03) !important; /* leve efeito de elevação */
                transition: all 0.2s !important;
            }

            /* Estilo para botões primários e de Download*/
            .stButton > button[kind="primary"],
            .stDownloadButton > button {
                background: linear-gradient(135deg, #FFF, #FFF) !important;
                border-radius: 10px !important;
                color: green !important;
            }
            .stButton > button[kind="primary"]:hover,
            .stDownloadButton > button:hover {
                background: linear-gradient(135deg, #FFF, #FFF) !important;
                color: #1bb50b !important; /* texto verde */
                border: 2px solid #1bb50b !important; /* borda verde */
                box-shadow: 0 2px 8px rgba(27,181,11,0.15) !important; /* sombra suave */
                transform: translateY(-2px) scale(1.03) !important; /* leve efeito de elevação */
                transition: all 0.2s !important;
            }

            /* Estilo para DataFrames */
            .stDataFrame {
                border-radius: 8px !important;
                border: 1px solid #e0e0e0 !important;
                overflow: hidden !important;
            }

            /* Estilo para as abas */
            .stTabs [data-baseweb="tab-list"] {
                gap: 10px;
            }
            
            /* Linha da aba ativa */
            div[data-baseweb="tab-highlight"] {
                background-color: #1bb50b; 
            }
            .stTabs [aria-selected="false"] {
                color: #000000 !important;
            }
            .stTabs [aria-selected="true"] {
                color: #1bb50b !important;
            }

            /* Estilo para file uploader */
            .stFileUploader {
                border: 2px dashed #e0e0e0 !important;
                border-radius: 8px !important;
                padding: 8px !important;
                text-align: center !important;
            }
        </style>
        """,unsafe_allow_html=True)
    st.markdown("<h1 style='text-align: center;'>Comparador de Arquivos e Textos da GNCP</h1>", unsafe_allow_html=True)

# Interface principal
    tab1, tab2, tab3 = st.tabs([ "Comparar Textos", "Comparar Documentos", "Comparar Planilhas Excel"])

    with tab1:
        col1, col2 = st.columns(2)

        # Inicializa contador de reset
        if "txt_reset" not in st.session_state:
            st.session_state.txt_reset = 0

        # Gera as chaves dinâmicas com base no contador
        txt1_key = f"txt1_input_{st.session_state.txt_reset}"
        txt2_key = f"txt2_input_{st.session_state.txt_reset}"

        with col1:
            st.text_area(
                "Texto Original",
                key=txt1_key,
                height=300,
                help="Digite ou cole o texto original para ser comparado"
            )

        with col2:
            st.text_area(
                "Texto Modificado",
                key=txt2_key,
                height=300,
                help="Digite ou cole o segundo texto para ser comparado"
            )

        # Verifica se os dois textos foram preenchidos
        txt1 = st.session_state.get(txt1_key, "")
        txt2 = st.session_state.get(txt2_key, "")

        if txt1 and txt2:
            btn_col1, btn_col2 = st.columns([1, 1])

            with btn_col1:
                comparar = st.button("Comparar Textos", type="primary")
            with btn_col2:
                apagar = st.button("Limpar Textos")

            if comparar:
                comparison = compare_texts(txt1, txt2)
                st.markdown(comparison, unsafe_allow_html=True)

                st.download_button(
                    label="Baixar Comparação",
                    data=comparison.encode("utf-8"),
                    file_name="Textos Comparados.html",
                    mime="text/html"
                )

            if apagar:
                st.session_state.txt_reset += 1
                st.rerun()

    with tab2:
        col1, col2 = st.columns(2)

        # Adicione um contador de reset no session_state
        if "arq_reset" not in st.session_state:
            st.session_state.arq_reset = 0

        with col1:
            st.session_state.arq1 = st.file_uploader(
                "Carregar Documento 1",
                type=["doc", "docx", "pdf", "txt", "csv"],
                accept_multiple_files=False,
                key=f"file1_input_{st.session_state.arq_reset}",
                help="Carregar primeiro arquivo para comparação"
            )
        with col2:
            st.session_state.arq2 = st.file_uploader(
                "Carregar Documento 2",
                type=["doc", "docx", "pdf", "txt", "csv"],
                accept_multiple_files=False,
                key=f"file2_input_{st.session_state.arq_reset}",
                help="Carregar segundo arquivo para comparação"
            )

        if st.session_state.arq1 and st.session_state.arq2:
            btn_col1, btn_col2 = st.columns([1, 1])
            with btn_col1:
                comparar = st.button("Comparar Documentos", key="comparar_docs", type="primary")
            with btn_col2:
                limpar = st.button("Limpar Uploads", key="limpar_docs")

            text1 = st.session_state.arq1.name.split('.')[-1].lower()
            text2 = st.session_state.arq2.name.split('.')[-1].lower()

            if comparar:
                with st.spinner("Fazendo comparações..."):
                    if text1 != text2:
                        st.error("Só é possível comparar arquivos do mesmo tipo/extensão!")
                    else:
                        result_doc, diff_doc, iguais = compare_docs(st.session_state.arq1, st.session_state.arq2)
                        if iguais:
                            st.info("Os arquivos são idênticos!")
                        elif result_doc and (diff_doc is None or not diff_doc.empty):
                            st.markdown(result_doc, unsafe_allow_html=True)

                            # Comparar tabelas se existirem
                            file_ext = text1  # já foi definido como a extensão
                            tables1 = tables2 = []

                            if file_ext == "pdf":
                                tables1 = extract_tables_from_pdf(st.session_state.arq1)
                                tables2 = extract_tables_from_pdf(st.session_state.arq2)
                            elif file_ext == "docx":
                                tables1 = extract_tables_from_docx(st.session_state.arq1)
                                tables2 = extract_tables_from_docx(st.session_state.arq2)

                            if tables1 or tables2:
                                st.markdown("---")
                                st.markdown("## Comparação de Tabelas Detectadas")
                                table_results = compare_table_list(tables1, tables2)
                                for idx, result in enumerate(table_results):
                                    display_excel_comparison(result, f"Tabela {idx+1}", st.session_state.arq1.name, st.session_state.arq2.name)
                            
                            st.download_button(
                                label="Baixar Comparação",
                                data=result_doc.encode("utf-8"),
                                file_name="Arquivos Comparados.html",
                                mime="text/html"
                            )
                        else:
                            st.error("Não foi possível comparar os documentos. Verifique os formatos.")

            elif limpar:
                st.session_state.arq1 = None
                st.session_state.arq2 = None
                st.session_state.arq_reset += 1  # incrementa para forçar reset dos file_uploaders
                st.rerun()

    with tab3:
        col1, col2 = st.columns(2)

        if "file_reset" not in st.session_state:
            st.session_state.file_reset = 0
        
        with col1:
            st.session_state.file1 = st.file_uploader(
                "Carregar Arquivo Excel 1", 
                type=["xlsx", "xls"],
                accept_multiple_files=False, 
                key=f"wb1_{st.session_state.file_reset}",
                help="Carregar primeira planilha para comparação"
            )
        with col2:
            st.session_state.file2 = st.file_uploader(
                "Carregar Arquivo Excel 2", 
                type=["xlsx", "xls"],
                accept_multiple_files=False, 
                key=f"wb2_{st.session_state.file_reset}",
                help="Carregar segunda planilha para comparação"
            )
            
        if st.session_state.file1 and st.session_state.file2:
            # Verificar se os arquivos são os mesmos
            if st.session_state.file1.name == st.session_state.file2.name:
                st.warning("Você carregou o mesmo arquivo duas vezes!")
            # Verificar se os arquivos são idênticos   
            elif excel_equal(st.session_state.file1, st.session_state.file2):
                st.success("Os arquivos são idênticos!")
            else:
                try:
                    xls1 = pd.ExcelFile(st.session_state.file1)
                    xls2 = pd.ExcelFile(st.session_state.file2)
                    all_sheets = sorted(set(xls1.sheet_names) | set(xls2.sheet_names))
                    
                    # Seleção de aba
                    if len(all_sheets) > 1:
                        selected_sheet = st.selectbox(
                            "Selecione a aba para comparar:",
                            options=all_sheets,
                            index=0,
                            help="Escolha qual aba/planilha deseja comparar"
                        )
                    else:
                        selected_sheet = all_sheets[0] if all_sheets else None
                        st.info(f"Comparando aba: **{selected_sheet}**")
                    
                    btn_col1, btn_col2 = st.columns([1, 1])
                    with btn_col1:
                        comparar = st.button("Comparar Planilhas", key="comparar_excel", type="primary")
                    with btn_col2:
                        limpar = st.button("Limpar Uploads", key="limpar_excel")
                        
                    if comparar:
                        with st.spinner("Comparando arquivos..."):
                            result = compare_excel(st.session_state.file1, st.session_state.file2, selected_sheet)
                            if result:
                                display_excel_comparison(result, selected_sheet, st.session_state.file1.name, st.session_state.file2.name)
                                
                    elif limpar:
                        st.session_state.file1 = None
                        st.session_state.file2 = None
                        st.session_state.file_reset += 1
                        st.rerun()

                except Exception as e:
                    st.error(f"Erro ao processar arquivos: {e}")
                    st.info("Verifique se os arquivos são válidos e não estão corrompidos.")

if __name__ == "__main__":
    main()